from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
GRAY_FILL = "F2F4F7"
BORDER = "D9E2EC"


def set_cell_fill(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_border(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
      borders = OxmlElement("w:tcBorders")
      tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
      tag = f"w:{edge}"
      element = borders.find(qn(tag))
      if element is None:
        element = OxmlElement(tag)
        borders.append(element)
      element.set(qn("w:val"), "single")
      element.set(qn("w:sz"), "6")
      element.set(qn("w:space"), "0")
      element.set(qn("w:color"), BORDER)


def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths[idx])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_border(cell)


def add_footer(section, label):
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run(label)
    run.font.name = "Calibri"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(90, 105, 96)


def configure_doc(title, subtitle):
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    add_footer(section, "RentEase submission document")

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.10

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_run = title_p.add_run(title)
    title_run.font.name = "Calibri"
    title_run.font.size = Pt(24)
    title_run.font.bold = True
    title_run.font.color.rgb = BLUE

    subtitle_p = doc.add_paragraph()
    subtitle_run = subtitle_p.add_run(subtitle)
    subtitle_run.font.name = "Calibri"
    subtitle_run.font.size = Pt(12)
    subtitle_run.font.color.rgb = RGBColor(90, 105, 96)

    doc.add_paragraph("Prepared for project submission")
    return doc


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths)
    hdr = table.rows[0].cells
    for idx, header in enumerate(headers):
        hdr[idx].text = header
        set_cell_fill(hdr[idx], GRAY_FILL)
        for paragraph in hdr[idx].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True

    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
    set_table_geometry(table, widths)
    return table


def build_prd():
    doc = configure_doc("RentEase Product Requirements Document", "Furniture and appliance rental platform")
    doc.add_heading("1. Product Overview", level=1)
    doc.add_paragraph(
        "RentEase is a rental platform where users register, login, browse furniture and appliances, add products to a rental cart, schedule delivery, manage active rentals, and request maintenance support."
    )
    doc.add_heading("2. Goals", level=1)
    add_bullets(doc, [
        "Require account registration or login before users can access catalog, rentals, checkout, support, or admin features.",
        "Provide a clean green and white interface with separate pages for major workflows.",
        "Support furniture and appliance catalog browsing, rental cart management, delivery scheduling, and support requests.",
        "Provide an admin dashboard for inventory, revenue, maintenance, delivery, and claims visibility.",
    ])
    doc.add_heading("3. User Roles", level=1)
    add_table(doc, ["Role", "Primary needs"], [
        ["Customer", "Register, login, browse catalog, create rentals, view rental history, and submit support requests."],
        ["Admin", "Monitor inventory, business metrics, maintenance requests, pending deliveries, and restock actions."],
    ], [1.6, 4.9])
    doc.add_heading("4. Functional Requirements", level=1)
    add_bullets(doc, [
        "Registration must create a new customer account and immediately authenticate the user.",
        "Login must authenticate registered users and store a session token for protected API access.",
        "Catalog must show furniture and appliance products with category filters.",
        "Cart must support add, remove, item count, monthly rent total, and security deposit total.",
        "Checkout must collect delivery date, city, location, and tenure before creating a rental.",
        "Rentals must display active and historical plans with status and next action.",
        "Support must create requests for repair, installation, pickup delay, and damage claim issues.",
        "Admin dashboard must show MRR, open requests, pending deliveries, damage claims, and product availability.",
    ])
    doc.add_heading("5. Acceptance Criteria", level=1)
    add_numbered(doc, [
        "A guest who opens the site is redirected to login or register.",
        "A registered user can login and browse catalog pages without server unreachable errors.",
        "Each top navigation item opens a separate route page.",
        "A selected product can be added to the cart and confirmed through checkout.",
        "The production build completes successfully and serves through the Express application.",
    ])
    doc.save("docs/RentEase_PRD.docx")


def build_technical():
    doc = configure_doc("RentEase Technical Documentation", "Architecture, database, APIs, stack, and implementation")
    doc.add_heading("1. Architecture", level=1)
    doc.add_paragraph(
        "RentEase is implemented as a React single page application served by a Node.js and Express.js backend. The frontend calls REST APIs for authentication, catalog, rentals, support, and admin data. Protected pages require a valid JWT token."
    )
    doc.add_heading("2. Technology Stack", level=1)
    add_table(doc, ["Layer", "Technology"], [
        ["Frontend", "HTML5, CSS3, JavaScript, React.js, React Router, Bootstrap, Vite"],
        ["Backend", "Node.js, Express.js, REST APIs, JWT authentication"],
        ["Database", "MongoDB-ready configuration through MONGO_URI with a local JSON preview store"],
        ["Deployment", "Vite production build served by Express from the dist folder"],
    ], [1.5, 5.0])
    doc.add_heading("3. API Endpoints", level=1)
    add_table(doc, ["Method", "Endpoint", "Purpose"], [
        ["POST", "/api/auth/register", "Create user account and return token"],
        ["POST", "/api/auth/login", "Authenticate user and return token"],
        ["GET", "/api/me", "Return current user profile"],
        ["GET", "/api/products", "Return protected catalog data"],
        ["GET", "/api/rentals", "Return user rental plans"],
        ["POST", "/api/rentals", "Create a rental from checkout"],
        ["POST", "/api/support", "Create a maintenance support request"],
        ["GET", "/api/admin/summary", "Return admin dashboard metrics"],
        ["POST", "/api/admin/restock", "Increase low-stock product availability"],
    ], [1.0, 2.0, 3.5])
    doc.add_heading("4. Data Model", level=1)
    add_bullets(doc, [
        "User: id, name, email, passwordHash, city, role.",
        "Product: id, name, category, rent, deposit, tenure, stock, image.",
        "Rental: id, email, product, tenure, status, nextAction, monthlyRent, city, location.",
        "Support request: id, email, product, issueType, details, status.",
    ])
    doc.add_heading("5. Implementation Notes", level=1)
    add_bullets(doc, [
        "All app pages except login and register are wrapped in a protected route.",
        "The frontend stores the JWT and user profile in browser local storage after login or registration.",
        "The backend validates bearer tokens before serving catalog, rental, support, and admin APIs.",
        "A local JSON store allows offline preview; MongoDB can be configured for deployment with MONGO_URI.",
    ])
    doc.add_heading("6. Deployment Steps", level=1)
    add_numbered(doc, [
        "Install dependencies with npm install.",
        "Set PORT, JWT_SECRET, and MONGO_URI in the environment.",
        "Run npm run build to create the production frontend.",
        "Run npm run server or npm start to serve the API and built React app.",
    ])
    doc.save("docs/RentEase_Technical_Documentation.docx")


if __name__ == "__main__":
    build_prd()
    build_technical()
